import streamlit as st
import json
import re
import time
import asyncio
import logging
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
import hashlib
from transformers import pipeline
from scholarly import scholarly
import os
from pdf2docx import Converter
import tempfile
from dotenv import load_dotenv

# Configuration and Constants
load_dotenv()
API_TIMEOUT = 30
MAX_RETRIES = 3
RETRY_DELAY = 2
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class Citation:
    text: str
    context: str
    line_number: int
    full_sentence: str
    citation_id: str
    suggested_citation: Optional[str] = None
    is_valid: bool = False
    replacement_info: Optional[Dict] = None
    scholar_data: Optional[Dict] = None

class EnhancedCitationExtractor:
    def __init__(self):
        self.ner_model = pipeline("ner", model="allenai/scibert_scivocab_uncased")
        self.citation_patterns = [
            r'\b[A-Z][a-z]+(?:\s+and\s+[A-Z][a-z]+)?\s*\(\d{4}\)',
            r'\b[A-Z][a-z]+(?:\s+et\s+al\.)?\s*\(\d{4}\)',
            r'\b[A-Z][a-z]+(?:\s+&\s+[A-Z][a-z]+)?\s*\(\d{4}\)',
            r'\([^)]+?,\s*\d{4}[^)]*?\)',
            r'\[[^\]]+?,\s*\d{4}[^\]]*?\]'
        ]

    def extract_citations(self, text: str) -> List[Citation]:
        citations = []
        seen_citations = {}
        
        for pattern in self.citation_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                citation_text = match.group()
                position = match.start()
                if citation_text not in seen_citations:
                    seen_citations[citation_text] = position
                    if self._is_valid_citation(citation_text):
                        citations.append(self._create_citation(
                            citation_text,
                            self._get_context(text, match.start(), match.end()),
                            self._get_line_number(text, match.start()),
                            self._get_full_sentence(text, match.start()),
                            position
                        ))

        ner_results = self.ner_model(text)
        for entity in ner_results:
            if entity['entity'] in ['B-REFERENCE', 'I-REFERENCE']:
                citation_text = text[entity['start']:entity['end']]
                if citation_text not in seen_citations and self._is_valid_citation(citation_text):
                    citations.append(self._create_citation(
                        citation_text,
                        self._get_context(text, entity['start'], entity['end']),
                        self._get_line_number(text, entity['start']),
                        self._get_full_sentence(text, entity['start']),
                        entity['start']
                    ))
                    seen_citations[citation_text] = entity['start']

        return sorted(citations, key=lambda x: x.line_number)

    def _create_citation(self, citation_text: str, context: str, line_num: int, sentence: str, position: int) -> Citation:
        return Citation(
            text=citation_text,
            context=context,
            line_number=line_num,
            full_sentence=sentence.strip(),
            citation_id=hashlib.md5(f"{citation_text}{context}{position}".encode()).hexdigest(),
            is_valid=True
        )

    def _get_context(self, text: str, start: int, end: int, window: int = 200) -> str:
        context_start = max(0, start - window)
        context_end = min(len(text), end + window)
        return text[context_start:context_end].strip()

    def _get_line_number(self, text: str, position: int) -> int:
        return text[:position].count('\n') + 1

    def _get_full_sentence(self, text: str, position: int) -> str:
        sentences = re.split(r'(?<=[.!?])\s+', text)
        current_pos = 0
        for sentence in sentences:
            sentence_length = len(sentence) + 1
            if current_pos <= position < current_pos + sentence_length:
                return sentence
            current_pos += sentence_length
        return ""

    def _is_valid_citation(self, citation_text: str) -> bool:
        year_match = re.search(r'\b(19|20)\d{2}\b', citation_text)
        if not year_match:
            return False
        year = int(year_match.group())
        current_year = datetime.now().year
        return 1900 <= year <= current_year and bool(re.search(r'[A-Z][a-z]+', citation_text))

class ScholarlyAPI:
    def __init__(self):
        self.cache = {}
        self._request_lock = asyncio.Lock()
        self._last_request = 0
        self.min_delay = 1.0

    async def search_citation(self, citation: Citation) -> Dict:
        if citation.citation_id in self.cache:
            return self.cache[citation.citation_id]

        async with self._request_lock:
            current_time = time.time()
            time_since_last = current_time - self._last_request
            if time_since_last < self.min_delay:
                await asyncio.sleep(self.min_delay - time_since_last)
            
            try:
                author_year = re.match(r'([A-Za-z\s&]+)\s*\((\d{4})\)', citation.text)
                if not author_year:
                    return None

                author, year = author_year.groups()
                search_query = f"author:{author} year:{year}"
                
                result = await asyncio.to_thread(
                    lambda: next(scholarly.search_pubs(search_query), None)
                )
                
                if result:
                    paper_data = {
                        'title': result.bib.get('title'),
                        'authors': result.bib.get('author', []),
                        'year': result.bib.get('year'),
                        'venue': result.bib.get('venue'),
                        'citations': result.citedby,
                        'url': result.bib.get('url')
                    }
                    self.cache[citation.citation_id] = paper_data
                    return paper_data

            except Exception as e:
                logger.error(f"Scholar API error: {str(e)}")
                return None
            finally:
                self._last_request = time.time()

class OpenAIAgent:
    def __init__(self, api_key=None, model="gpt-3.5-turbo"):
        self.client = OpenAI(api_key=api_key)
        self.model = model
        self.cache = {}
        self.requests_per_minute = 50
        self.last_request_time = time.time()
        self.request_times = []
        self._request_lock = asyncio.Lock()

    async def suggest_updated_citation(self, citation: Citation, scholar_data: Optional[Dict] = None) -> Dict:
        cache_key = citation.citation_id
        if cache_key in self.cache:
            return self.cache[cache_key]

        async with self._request_lock:
            prompt = self._create_update_prompt(citation, scholar_data)
            try:
                response = await self._call_api_async(prompt)
                try:
                    result = json.loads(response)
                    if self._validate_suggestion(result):
                        self.cache[cache_key] = result
                        return result
                except json.JSONDecodeError:
                    logger.error(f"Invalid JSON response for citation {citation.text}")
                return None
            except Exception as e:
                logger.error(f"OpenAI API error: {str(e)}")
                return None

    def _validate_suggestion(self, suggestion: Dict) -> bool:
        required_fields = ['suggested_paper', 'relevance_explanation']
        paper_fields = ['authors', 'year', 'title', 'journal', 'volume', 'issue', 'pages', 'doi']
        return all(field in suggestion for field in required_fields) and all(field in suggestion['suggested_paper'] for field in paper_fields)

    def _create_update_prompt(self, citation: Citation, scholar_data: Optional[Dict]) -> str:
        return f"""
Citation: {citation.text}
Context: {citation.context}
Original Paper Data: {json.dumps(scholar_data) if scholar_data else 'Not available'}

Please suggest a recent (2020-2024) relevant paper to update this citation. Return in JSON format:
{{
    "suggested_paper": {{
        "authors": ["Surname1, F. M.", "Surname2, F. M."],
        "year": YYYY,
        "title": "Title in sentence case",
        "journal": "Journal Title in Title Case",
        "volume": "Vol",
        "issue": "Issue",
        "pages": "start-end",
        "doi": "doi-string"
    }},
    "relevance_explanation": "Brief explanation of why this paper is a suitable update"
}}
"""

    async def _call_api_async(self, prompt: str) -> str:
        current_time = time.time()
        self.request_times = [t for t in self.request_times if current_time - t < 60]
        if len(self.request_times) >= self.requests_per_minute:
            await asyncio.sleep(60 - (current_time - self.request_times[0]))

        for attempt in range(MAX_RETRIES):
            try:
                response = await asyncio.to_thread(
                    self.client.chat.completions.create,
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "You are a citation expert specializing in finding recent relevant academic papers."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,
                    max_tokens=500
                )
                self.request_times.append(time.time())
                return response.choices[0].message.content
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(RETRY_DELAY * (attempt + 1))
                    continue
                raise e

class CitationProcessor:
    def __init__(self, api_key: str):
        self.extractor = EnhancedCitationExtractor()
        self.scholar_api = ScholarlyAPI()
        self.openai_agent = OpenAIAgent(api_key=api_key)

    async def process_document(self, text: str) -> Tuple[str, List[Dict], List[str]]:
        citations = self.extractor.extract_citations(text)
        updated_text = text
        citation_changes = []
        citation_reference_pairs = {}
        citation_positions = []

        for citation in citations:
            scholar_data = await self.scholar_api.search_citation(citation)
            update_suggestion = await self.openai_agent.suggest_updated_citation(
                citation, scholar_data
            )
            
            if update_suggestion and 'suggested_paper' in update_suggestion:
                paper = update_suggestion['suggested_paper']
                new_citation = self._format_citation(paper['authors'], paper['year'])
                reference = self._format_reference(paper)
                
                start_pos = text.find(citation.text)
                if start_pos != -1:
                    citation_positions.append((start_pos, citation.text, new_citation))
                    citation_reference_pairs[new_citation] = {
                        'reference': reference,
                        'change': {
                            'original': citation.text,
                            'updated': new_citation,
                            'context': citation.context,
                            'explanation': update_suggestion['relevance_explanation']
                        }
                    }

        citation_positions.sort(reverse=True)
        
        for _, old_citation, new_citation in citation_positions:
            if old_citation in updated_text:
                updated_text = updated_text.replace(old_citation, new_citation)
                if citation_reference_pairs[new_citation]['change'] not in citation_changes:
                    citation_changes.append(citation_reference_pairs[new_citation]['change'])

        used_citations = {citation for citation in citation_reference_pairs.keys() 
                         if updated_text.count(citation) > 0}
        references = [citation_reference_pairs[citation]['reference'] 
                     for citation in used_citations]

        if len(used_citations) != len(references):
            logger.error("Citation-reference count mismatch")
            return text, [], []

        citation_counts = {citation: updated_text.count(citation) for citation in used_citations}
        if any(count > 1 for count in citation_counts.values()):
            logger.error("Duplicate citations detected")
            return text, [], []

        if any(citation not in citation_reference_pairs for citation in used_citations):
            logger.error("Missing reference detected")
            return text, [], []

        return updated_text, citation_changes, sorted(references)

    def _format_citation(self, authors: List[str], year: str) -> str:
        if len(authors) > 2:
            return f"{authors[0].split(',')[0]} et al. ({year})"
        elif len(authors) == 2:
            return f"{authors[0].split(',')[0]} & {authors[1].split(',')[0]} ({year})"
        return f"{authors[0].split(',')[0]} ({year})"

    def _format_reference(self, paper: Dict) -> str:
        return (
            f"{', '.join(paper['authors'])}. ({paper['year']}). {paper['title']}. "
            f"{paper['journal']}, {paper['volume']}({paper['issue']}), {paper['pages']}. "
            f"https://doi.org/{paper['doi']}"
        )

def create_word_document(original_text: str, updated_text: str, 
                        citation_changes: List[Dict], references: List[str]) -> BytesIO:
    doc = Document()
    
    style = doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    title = doc.add_heading('Updated Academic Document', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Original Text', level=2)
    para = doc.add_paragraph(original_text)
    para.style = 'CustomNormal'
    
    doc.add_heading('Updated Text', level=2)
    para = doc.add_paragraph(updated_text)
    para.style = 'CustomNormal'
    
    if citation_changes:
        doc.add_heading('Citation Updates', level=2)
        for change in citation_changes:
            para = doc.add_paragraph()
            para.style = 'CustomNormal'
            original = para.add_run(f"Original: {change['original']}\n")
            original.font.color.rgb = RGBColor(192, 0, 0)
            updated = para.add_run(f"Updated: {change['updated']}\n")
            updated.font.color.rgb = RGBColor(0, 100, 0)
            para.add_run(f"Context: {change['context']}\n")
            para.add_run(f"Reason for update: {change['explanation']}\n\n")
    
    if references:
        doc.add_heading('References', level=2)
        for ref in references:
            para = doc.add_paragraph(ref)
            para.style = 'CustomNormal'
            para.paragraph_format.first_line_indent = Inches(-0.5)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(12)
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

def extract_text_from_pdf(pdf_file) -> str:
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
        tmp_pdf.write(pdf_file.getvalue())
        tmp_pdf_path = tmp_pdf.name
        
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
        tmp_docx_path = tmp_docx.name
        
    try:
        converter = Converter(tmp_pdf_path)
        converter.convert(tmp_docx_path)
        converter.close()
        
        doc = Document(tmp_docx_path)
        text_content = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        return text_content
    finally:
        os.unlink(tmp_pdf_path)
        os.unlink(tmp_docx_path)

def main():
    st.set_page_config(
        page_title="PyBib | Academic Reference Manager",
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Apply custom styling
    st.markdown("""
        <style>
        /* Global Styles */
        .main {
            padding: 2rem;
            max-width: 1200px;
            margin: 0 auto;
            background-color: #FAFBFC;
        }
        
        /* Custom Fonts */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');
        
        /* Button Styles */
        .stButton > button {
            width: 100%;
            background: linear-gradient(135deg, #4F46E5 0%, #3B82F6 100%);
            color: white;
            border-radius: 8px;
            padding: 0.75rem 1.5rem;
            font-weight: 600;
            font-family: 'Inter', sans-serif;
            border: none;
            transition: all 0.3s ease;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        }
        
        .stButton > button:hover {
            background: linear-gradient(135deg, #4338CA 0%, #2563EB 100%);
            transform: translateY(-1px);
            box-shadow: 0 6px 8px -1px rgba(0, 0, 0, 0.1);
        }
        
        /* Layout Container */
        .css-1v0mbdj.e115fcil1 {
            max-width: 1200px;
            margin: 0 auto;
            background-color: #FFFFFF;
            border-radius: 12px;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1);
        }
        
        /* Cards and Containers */
        .report-box {
            border: 1px solid #E5E7EB;
            border-radius: 12px;
            padding: 2rem;
            margin: 1.5rem 0;
            background-color: #FFFFFF;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            transition: transform 0.2s ease;
        }
        
        .report-box:hover {
            transform: translateY(-2px);
        }
        
        .metric-card {
            background: linear-gradient(135deg, #FFFFFF 0%, #F9FAFB 100%);
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            border: 1px solid #E5E7EB;
            transition: all 0.3s ease;
        }
        
        .metric-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 8px -1px rgba(0, 0, 0, 0.1);
        }
        
        .metric-card h3 {
            color: #4B5563;
            font-size: 1rem;
            font-weight: 600;
            margin-bottom: 0.5rem;
        }
        
        .metric-card h2 {
            color: #1F2937;
            font-size: 2rem;
            font-weight: 700;
        }
        
        /* Citation Cards */
        .citation-card {
            border-left: 4px solid #4F46E5;
            padding: 1.5rem;
            margin: 1rem 0;
            background: linear-gradient(to right, #F9FAFB, #FFFFFF);
            border-radius: 0 12px 12px 0;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            transition: all 0.2s ease;
        }
        
        .citation-card:hover {
            transform: translateX(4px);
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.07);
        }
        
        .citation-card h4 {
            color: #4F46E5;
            font-weight: 600;
            margin-bottom: 1rem;
        }
        
        /* File Uploader */
        .uploadedFile {
            border: 2px dashed #E5E7EB;
            border-radius: 12px;
            padding: 2rem;
            text-align: center;
            background-color: #F9FAFB;
            transition: all 0.3s ease;
        }
        
        .uploadedFile:hover {
            border-color: #4F46E5;
            background-color: #F3F4F6;
        }
        
        /* Reference Items */
        .reference-item {
            padding: 1rem;
            margin: 0.5rem 0;
            border-radius: 8px;
            background-color: #F9FAFB;
            border: 1px solid #E5E7EB;
            transition: all 0.2s ease;
        }
        
        .reference-item:hover {
            background-color: #F3F4F6;
            transform: translateX(2px);
        }
        
        /* Text Areas */
        .stTextArea > div > div {
            border-radius: 8px;
            border-color: #E5E7EB;
            background-color: #F9FAFB;
        }
        </style>
    """, unsafe_allow_html=True)

    # Header
    st.markdown("""
        <div style='text-align: center; padding: 2rem 0; background: linear-gradient(135deg, #F9FAFB 0%, #F3F4F6 100%); border-radius: 12px; margin-bottom: 2rem;'>
            <h1 style='font-family: Inter, sans-serif; font-size: 2.5rem; font-weight: 700; 
                      background: linear-gradient(135deg, #4F46E5 0%, #3B82F6 100%);
                      -webkit-background-clip: text; -webkit-text-fill-color: transparent;
                      margin-bottom: 1rem;'>
                Citation Pro
            </h1>
            <p style='font-family: Inter, sans-serif; color: #4B5563; font-size: 1.2rem; 
                      max-width: 600px; margin: 0 auto;'>
                Transform your academic references with AI-powered citation management
            </p>
        </div>
    """, unsafe_allow_html=True)

    # Sidebar Configuration
    with st.sidebar:
        st.markdown("""
            <div style='padding: 1.5rem; background: linear-gradient(135deg, #4F46E5 0%, #3B82F6 100%); 
                        border-radius: 12px; margin-bottom: 1.5rem;'>
                <h3 style='color: white; font-family: Inter, sans-serif; font-weight: 600; 
                           margin-bottom: 0.5rem;'>🚀 Citation Pro</h3>
                <p style='color: #E5E7EB; font-size: 0.9rem;'>Powered by Advanced AI</p>
            </div>
        """, unsafe_allow_html=True)
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.markdown("""
                <div style='padding: 1rem; background-color: #FEF2F2; border: 1px solid #FCA5A5;
                            border-radius: 8px; margin-bottom: 1rem;'>
                    <h4 style='color: #991B1B; margin-bottom: 0.5rem;'>⚠️ API Key Missing</h4>
                    <p style='color: #7F1D1D; font-size: 0.9rem;'>Please configure your OpenAI API key.</p>
                </div>
            """, unsafe_allow_html=True)
            return

        st.markdown("""
            <div style='padding: 1rem; background-color: #F3F4F6; border-radius: 8px; margin-bottom: 1rem;'>
                <h4 style='color: #1F2937; margin-bottom: 0.5rem;'>🔑 API Status</h4>
                <p style='color: #059669; font-size: 0.9rem;'>✓ Connected and ready</p>
            </div>
        """, unsafe_allow_html=True)

    # Main Content
    st.markdown("""
        <div class='uploadedFile'>
            <h3 style='color: #4F46E5; margin-bottom: 1rem;'>📄 Document Upload</h3>
            <p style='color: #6B7280;'>Supported formats: PDF, Word, and Text files</p>
        </div>
    """, unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "",
        type=["txt", "docx", "pdf"],
        help="Upload your academic document for citation analysis"
    )

    if uploaded_file:
        try:
            with st.spinner("📑 Reading document..."):
                if uploaded_file.type == "text/plain":
                    text_content = uploaded_file.getvalue().decode('utf-8')
                elif uploaded_file.type == "application/pdf":
                    text_content = extract_text_from_pdf(uploaded_file)
                else:
                    doc = Document(uploaded_file)
                    text_content = '\n'.join(paragraph.text for paragraph in doc.paragraphs)

            col1, col2, col3 = st.columns([1,2,1])
            with col2:
                process_button = st.button("🔄 Process Citations", use_container_width=True)

            if process_button:
                with st.spinner("🔍 Analyzing citations..."):
                    processor = CitationProcessor(api_key)
                    updated_text, changes, references = asyncio.run(
                        processor.process_document(text_content)
                    )

                    if not changes or not references:
                        st.error("⚠️ No citations found or error in processing. Please check your document format.")
                        return

                    citation_count = len([change for change in changes if change['updated'] in updated_text])
                    reference_count = len(references)

                    # Create Word document
                    doc_stream = create_word_document(text_content, updated_text, changes, references)

                    # Display Metrics
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(
                            f"""<div class='metric-card'>
                                <h3>📝 Citations Updated</h3>
                                <h2>{citation_count}</h2>
                            </div>""", 
                            unsafe_allow_html=True
                        )
                    with col2:
                        st.markdown(
                            f"""<div class='metric-card'>
                                <h3>📚 References Added</h3>
                                <h2>{reference_count}</h2>
                            </div>""", 
                            unsafe_allow_html=True
                        )

                    # Document Comparison
                    st.markdown("### 📑 Document Comparison")
                    tab1, tab2 = st.tabs(["Original Text", "Updated Text"])
                    with tab1:
                        st.text_area("", text_content, height=300, disabled=True)
                    with tab2:
                        st.text_area("", updated_text, height=300, disabled=True)

                    # Citation Updates
                    if changes:
                        st.markdown("### 🔄 Citation Updates")
                        for idx, change in enumerate(changes, 1):
                            if change['updated'] in updated_text:
                                st.markdown(
                                    f"""<div class='citation-card'>
                                        <h4>Update #{idx}</h4>
                                        <p><strong>Original:</strong> {change['original']}</p>
                                        <p><strong>Updated:</strong> {change['updated']}</p>
                                        <p><strong>Context:</strong> {change['context']}</p>
                                        <p><strong>Reason:</strong> {change['explanation']}</p>
                                    </div>""",
                                    unsafe_allow_html=True
                                )

                    # References
                    if references:
                        st.markdown("### 📚 References")
                        for idx, ref in enumerate(references, 1):
                            st.markdown(
                                f"""<div class='reference-item'>
                                    {idx}. {ref}
                                </div>""",
                                unsafe_allow_html=True
                            )

                    # Download Options
                    st.markdown("### 💾 Export Results")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            "📄 Download Word Document",
                            data=doc_stream,
                            file_name=f"citation_pro_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    with col2:
                        # Generate text report
                        text_report = f"""CITATION PRO - ACADEMIC DOCUMENT REPORT
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

ANALYSIS SUMMARY:
- Total Citations: {citation_count}
- Total References: {reference_count}

ORIGINAL TEXT:
{text_content}

UPDATED TEXT:
{updated_text}

CITATION CHANGES:
{'-' * 50}
{"".join(f'''
Update #{i+1}
Original: {change["original"]}
Updated: {change["updated"]}
Context: {change["context"]}
Reason: {change["explanation"]}
{'-' * 50}''' for i, change in enumerate(changes) if change['updated'] in updated_text)}

REFERENCES:
{'-' * 50}
{"".join(f'{i+1}. {ref}\n' for i, ref in enumerate(references))}"""

                        st.download_button(
                            "📝 Download Report",
                            data=text_report,
                            file_name=f"citation_report_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )

        except Exception as e:
            logger.error(f"Processing error: {str(e)}", exc_info=True)
            st.error(f"❌ An error occurred: {str(e)}")
            if "PDF" in str(e):
                st.info("📋 For PDFs, ensure the document contains selectable text.")
            else:
                st.info("📋 Please check your document format and try again.")

if __name__ == "__main__":
    main()